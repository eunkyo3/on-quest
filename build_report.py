# -*- coding: utf-8 -*-
"""On-Quest 최종보고서 생성 (python-docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY = "맑은 고딕"
MONO = "Consolas"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)

doc = Document()

def style_kfont(style, name=BODY):
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rfonts.set(qn(a), name)

for sn in ('Normal', 'Title', 'Heading 1', 'Heading 2', 'Heading 3'):
    try:
        style_kfont(doc.styles[sn])
    except KeyError:
        pass
doc.styles['Normal'].font.size = Pt(10.5)

for s in doc.sections:
    s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)
    s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)

def kfont_run(run, name=BODY, size=None, bold=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rfonts.set(qn(a), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color is not None: run.font.color.rgb = color

def h(text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text); kfont_run(r, BODY, color=ACCENT)
    return p

def para(text="", bold=False, size=10.5, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text); kfont_run(r, BODY, size, bold); r.font.italic = italic
    return p

def bullets(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        if isinstance(it, tuple):
            lead, rest = it
            r = p.add_run(lead); kfont_run(r, BODY, 10.5, True)
            r2 = p.add_run(rest); kfont_run(r2, BODY, 10.5, False)
        else:
            r = p.add_run(it); kfont_run(r, BODY, 10.5, False)

def mono_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8); p.paragraph_format.space_before = Pt(4)
    for i, ln in enumerate(lines):
        run = p.add_run(("" if i == 0 else "\n") + ln); kfont_run(run, MONO, 8.5)

def shade(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hexcolor)
    tcpr.append(sh)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(htext)
        kfont_run(r, BODY, 9.5, True, RGBColor(0xFF,0xFF,0xFF)); shade(hdr[i], '1F4E79')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val)); kfont_run(r, BODY, 9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def pagebreak():
    doc.add_page_break()

# ===== 표지 =====
for _ in range(3): para()
p = para(align=WD_ALIGN_PARAGRAPH.CENTER); r = p.add_run("2026-1학기 PBL 최종보고서"); kfont_run(r, BODY, 14, True, RGBColor(0x80,0x80,0x80))
para()
p = para(align=WD_ALIGN_PARAGRAPH.CENTER); r = p.add_run("On-Quest"); kfont_run(r, BODY, 34, True, ACCENT)
p = para(align=WD_ALIGN_PARAGRAPH.CENTER); r = p.add_run("게임형(퀘스트형) 신입사원 온보딩 자동화 플랫폼"); kfont_run(r, BODY, 14, True)
for _ in range(6): para()
info = [("프로젝트 형태", "개인 프로젝트"), ("재직 기업", "애니셀"),
        ("발표자", "정은교 (학번 2025800216)"), ("기업현장교사", "문용근"),
        ("지도교수", "김현우"), ("제출일", "2026년 6월")]
ti = doc.add_table(rows=0, cols=2); ti.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in info:
    c = ti.add_row().cells
    c[0].text = ""; rr = c[0].paragraphs[0].add_run(k); kfont_run(rr, BODY, 11, True)
    c[1].text = ""; rr = c[1].paragraphs[0].add_run(v); kfont_run(rr, BODY, 11, False)
    c[0].width = Inches(1.8); c[1].width = Inches(3.5)
pagebreak()

# ===== 목차 =====
h("목차", 1)
toc = ["1. 개요","   1.1 프로젝트 소개 및 목적","   1.2 주안점 (Key Point)","   1.3 주요 기술",
 "2. 요구사항 명세","   2.1 목적 및 범위","   2.2 기능적 요구사항","   2.3 비기능적 요구사항","   2.4 시스템 요구사항","   2.5 기타 요구사항 (비용·품질)",
 "3. 설계","   3.1 시스템 구조도","   3.2 데이터 모델 (ERD)","   3.3 주요 시퀀스 (증빙 제출·알림 흐름)","   3.4 퀘스트 상태 다이어그램","   3.5 UI 콘셉트",
 "4. 시스템 구현","   4.1 구현 환경","   4.2 주요 기능 구현 및 분석","   4.3 핵심 알고리즘 및 보안 설계","   4.4 기존 SW 대비 차별성·창의성",
 "5. 검증 (기능·성능 평가)","   5.1 기능 동작 체크리스트","   5.2 품질·성능 지표",
 "6. 소감 및 제언","7. 참고사항","   7.1 용어/약어","   7.2 참고문헌","부록. 실행 방법 및 소스 구조"]
for line in toc:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(line); kfont_run(r, BODY, 11, line[0].isdigit() or line.startswith("부록"))
pagebreak()

# ===== 1. 개요 =====
h("1. 개요", 1)
h("1.1 프로젝트 소개 및 목적", 2)
para("신입사원은 입사 직후 '무엇을 해야 할지' 명확히 알지 못해 혼란을 겪는 경우가 많다. 동시에 사수(선임)는 자신의 업무로 바빠 신입을 충분히 챙기지 못하고, 그 결과 신입은 업무가 주어지기 전까지 대기하거나 방치되어 조직 적응이 늦어진다.")
para("On-Quest는 이러한 '온보딩 공백'을 해소하기 위한 플랫폼이다. 회사가 신입에게 부여해야 할 할 일을 게임의 '퀘스트(Quest)'처럼 정의하고, 부여 → 수행 → 인증샷 제출 → 검토 → 피드백/완료에 이르는 전 과정을 하나의 워크플로우로 자동화한다. 또한 n8n과 Slack을 연동해 각 단계에서 실시간 알림을 발송함으로써, 관리자와 신입사원이 별도의 채근 없이도 온보딩 사이클을 진행할 수 있게 한다.")
para("최종 목적은 '신입사원이 입사하자마자 스스로 해야 할 일을 확인하고 빠르게 조직에 스며들 수 있도록 돕는 것'이며, 관리자의 온보딩 관리 부담을 동시에 경감하는 데 있다.")
h("1.2 주안점 (Key Point)", 2)
bullets([
 ("게임화(Gamification): ", "할 일을 단순 체크리스트가 아닌 상태(Status)를 가진 '퀘스트' 워크플로우로 모델링하여 진행 동기를 부여한다."),
 ("시스템 통합: ", "백엔드(NestJS) ↔ 자동화 툴(n8n) ↔ 메신저(Slack)를 유기적으로 연결하는 이벤트 기반 연동을 구현한다."),
 ("보안·신뢰성: ", "웹훅 HMAC 서명, JWT(HttpOnly 쿠키)·토큰 폐기, 회사별 데이터 격리, 감사 로그 등 실무 수준의 보안을 적용한다."),
 ("멀티테넌시: ", "회사코드 단위로 데이터를 완전히 분리하여 다수 기업이 한 인스턴스를 공유해도 안전하도록 설계한다."),
])
h("1.3 주요 기술", 2)
table(["구분", "기술"],
 [["언어", "TypeScript (프론트엔드·백엔드 공통)"],
  ["프론트엔드", "React 18 (Vite), Zustand(상태관리), React Router, Axios"],
  ["백엔드", "NestJS 10, Passport-JWT, class-validator, @nestjs/schedule, @nestjs/throttler"],
  ["데이터베이스 / ORM", "PostgreSQL 16 / Prisma"],
  ["자동화·연동", "n8n(워크플로우 자동화), Slack API, HMAC-SHA256 서명"],
  ["인프라·도구", "Docker Compose, Nginx, GitHub Actions(CI), Jest(테스트)"]],
 widths=[1.8, 4.6])
pagebreak()

# ===== 2. 요구사항 =====
h("2. 요구사항 명세", 1)
h("2.1 목적 및 범위", 2)
para("목적: 신입사원 온보딩 과제를 '부여–수행–검토–피드백'의 닫힌 루프로 자동화하고, 각 단계의 이벤트를 Slack으로 실시간 통지하여 온보딩 적응 속도를 높인다.")
bullets([
 ("범위(In-Scope): ", "회원/권한 관리, 퀘스트 발행(단건·CSV 일괄), 수행/제출, 검토(승인·반려·재개봉), 증빙 파일 관리, 마감 알림, 진행 통계, 감사 로그."),
 ("범위 외(Out-of-Scope): ", "급여·인사 시스템 연동, 모바일 네이티브 앱, 실시간 채팅(메신저 연동은 알림 수준)."),
])
h("2.2 기능적 요구사항", 2)
table(["ID", "역할", "기능", "설명"],
 [["FR-01","공통","회원가입/로그인","회사코드 기반 가입, JWT 인증. 회사 첫 가입자는 슈퍼관리자로 자동 지정"],
  ["FR-02","슈퍼관리자","구성원·권한 관리","구성원 역할 변경, 소유권(슈퍼관리자) 이양, 감사 로그 조회"],
  ["FR-03","관리자","퀘스트 발행","단건 발행 및 CSV 일괄 발행(최대 100건), 마감기한·담당자 지정"],
  ["FR-04","관리자","검토","제출 증빙 확인 후 승인(완료) 또는 피드백과 함께 반려, 거부 건 재개봉"],
  ["FR-05","신입사원","퀘스트 수행","할당 목록 확인, 착수, 수행 거부(사유), 인증샷(증빙) 제출·재제출"],
  ["FR-06","공통","증빙 파일","이미지/PDF 업로드, 미리보기·다운로드, Slack 외부 공유(서명 링크)"],
  ["FR-07","시스템","실시간 알림","발행·제출·검토·마감 임박/초과 시 n8n→Slack 자동 알림"],
  ["FR-08","관리자","진행 현황","진행률 대시보드, 담당자별 통계, 목록 CSV 내보내기"]],
 widths=[0.6,1.0,1.4,3.4])
h("2.3 비기능적 요구사항", 2)
table(["ID", "항목", "요구사항"],
 [["NFR-01","성능","Slack 실시간 알림은 이벤트 발생 후 3초 이내 전달(웹훅 비동기, 2.5s 타임아웃)"],
  ["NFR-02","보안","비밀번호 bcrypt 해시, JWT 서명, 모든 웹훅 페이로드 HMAC-SHA256 서명·검증"],
  ["NFR-03","보안","refresh 토큰 HttpOnly 쿠키 보관, 로그아웃 시 토큰 즉시 폐기(token versioning)"],
  ["NFR-04","보안","회사코드 기준 테넌트 격리, 역할 기반 접근제어(RBAC), 인증 시도 레이트리밋"],
  ["NFR-05","무결성","동시 요청에도 퀘스트 상태가 꼬이지 않도록 원자적 상태 전이 보장"],
  ["NFR-06","가용성","컨테이너 헬스체크 및 DB 프로브 타임아웃으로 장애 감지"],
  ["NFR-07","유지보수","TypeScript 정적 타입, ESLint, 자동화 테스트(CI) 통과 기준 유지"]],
 widths=[0.7,1.0,4.7])
h("2.4 시스템 요구사항", 2)
table(["구분", "요구사항"],
 [["서버 런타임","Node.js 20 이상"],["데이터베이스","PostgreSQL 16"],
  ["자동화 엔진","n8n (Docker 컨테이너)"],["실행 환경","Docker / Docker Compose"],
  ["클라이언트","최신 웹 브라우저(Chrome 등), 화면 폭 1024px 이상 권장"],
  ["외부 연동","Slack 워크스페이스(Incoming Webhook 또는 Bot)"]],
 widths=[1.8,4.6])
h("2.5 기타 요구사항 (비용·품질)", 2)
bullets([
 ("비용: ", "오픈소스 스택(NestJS·React·PostgreSQL·n8n)과 셀프호스팅 기반으로 라이선스 비용 없음. 인증샷은 외부 스토리지(S3) 대신 DB에 저장하여 운영 비용·외부 의존성을 최소화."),
 ("품질: ", "백엔드 단위/통합 테스트와 GitHub Actions CI(타입체크·린트·테스트·빌드)로 회귀를 방지하고 코드 일관성을 유지."),
])
pagebreak()

# ===== 3. 설계 =====
h("3. 설계", 1)
h("3.1 시스템 구조도", 2)
para("프론트엔드·백엔드·데이터베이스·자동화 엔진을 모두 Docker Compose로 구성하며, 각 컨테이너에 헬스체크를 둔다. 프론트엔드는 Nginx가 정적 파일을 서빙하고 /api 요청을 백엔드로 프록시한다.")
mono_block([
 "┌─────────────┐    HTTPS/REST    ┌──────────────────────┐",
 "│  Browser    │ ───────────────▶ │  Frontend (Nginx)    │",
 "│  (React)    │ ◀─────────────── │  React + Vite build  │",
 "└─────────────┘                  └──────────┬───────────┘",
 "                                  /api proxy │",
 "                                             ▼",
 "                                  ┌──────────────────────┐",
 "                       JWT/RBAC   │  Backend (NestJS)    │",
 "                                  │  auth·quest·audit·   │",
 "                                  │  automation·health   │",
 "                                  └───┬───────────────┬──┘",
 "                          Prisma ORM  │               │ HMAC 서명 webhook",
 "                                      ▼               ▼",
 "                          ┌──────────────┐    ┌──────────┐   Slack API",
 "                          │ PostgreSQL16 │    │   n8n    │ ───────────▶ Slack",
 "                          └──────────────┘    └──────────┘   (실시간 알림)",
])
h("3.2 데이터 모델 (ERD)", 2)
para("핵심 엔티티는 User, Quest, AuditLog 3개이며, Slack 멤버 ID와 회사코드를 통해 느슨하게 연결된다.")
table(["엔티티", "주요 필드", "설명"],
 [["User","email, passwordHash, name, slackMemberId, companyCode, role, tokenVersion",
   "구성원. (email, companyCode)·(slackMemberId, companyCode) 유니크. 회사당 슈퍼관리자 1명"],
  ["Quest","id, title, description, deadline, status, feedback, declineReason, proofData(BLOB), proofMimeType, assigneeId, reviewerId, companyCode, 알림플래그",
   "퀘스트. status는 0~5(대기/착수/검토대기/완료/반려/거부됨). 증빙은 BLOB로 저장"],
  ["AuditLog","companyCode, actorId, actorName, action, targetType, targetId, detail, createdAt",
   "권한변경·검토·삭제·이양 등 주요 행위 이력"]],
 widths=[0.9,2.7,2.8])
h("3.3 주요 시퀀스 (증빙 제출·알림 흐름)", 2)
mono_block([
 "신입사원        Frontend        Backend(NestJS)        DB        n8n        Slack",
 "  │  인증샷 제출   │                 │                 │          │           │",
 "  │ ─────────────▶ │  POST /proof    │                 │          │           │",
 "  │               │ ──────────────▶ │ 상태검증·원자적  │          │           │",
 "  │               │                 │ 업데이트(BLOB)   │          │           │",
 "  │               │                 │ ───────────────▶ │          │           │",
 "  │               │                 │   HMAC 서명 webhook(proof_uploaded)     │",
 "  │               │                 │ ───────────────────────────▶ │ 서명검증  │",
 "  │               │                 │                 │          │ ─────────▶ │ 알림",
 "  │               │  200 OK         │                 │          │           │",
 "  │ ◀───────────── │ ◀────────────── │                 │          │           │",
])
h("3.4 퀘스트 상태 다이어그램", 2)
mono_block([
 "             발행                착수                제출",
 "  (관리자) ───────▶ [대기] ───────────▶ [착수] ───────────▶ [검토대기]",
 "                     │                   │                     │",
 "             (사원)거부│           (사원)거부│            (관리자)검토",
 "                     ▼                   ▼               ┌─────┴─────┐",
 "                  [거부됨] ◀─────────────┘          승인 │           │ 반려",
 "                     │                                ▼           ▼",
 "             (관리자)재개봉 ───▶ [대기]            [완료]      [반려]",
 "                                                              │",
 "                                              재제출 ◀────────┘",
])
para("상태 전이는 모두 서버에서 '기대 상태일 때만 갱신되는' 원자적 연산으로 처리하여, 동시 요청이 들어와도 중복 처리나 상태 꼬임이 발생하지 않는다(검증 시 409 Conflict 반환).", size=9.5, italic=True)
h("3.5 UI 콘셉트", 2)
bullets([
 ("역할별 대시보드: ", "신입사원은 '내 퀘스트', 관리자는 '퀘스트 관리', 슈퍼관리자는 '사용자 관리'를 기본 화면으로 제공."),
 ("상태 가시화: ", "퀘스트마다 색상 배지로 상태(대기/착수/검토대기/완료/반려/거부)를 직관적으로 표시."),
 ("진행률 대시보드: ", "전체/담당자별 완료율을 한눈에 보여주고, 목록은 CSV로 내보낼 수 있음."),
 ("증빙 미리보기: ", "이미지 증빙은 모달로 즉시 미리보기, 그 외 형식은 안전하게 다운로드."),
])
para("[삽입 위치] 로그인 / 관리자 대시보드 / 신입 대시보드 / 퀘스트 상세 화면 캡처", size=9.5, italic=True)
pagebreak()

# ===== 4. 구현 =====
h("4. 시스템 구현", 1)
h("4.1 구현 환경", 2)
table(["구분", "내용"],
 [["OS / 런타임","Windows 10 / Node.js 20, Docker Desktop"],
  ["백엔드","NestJS 10, Prisma 5, Passport-JWT, class-validator, @nestjs/schedule·throttler"],
  ["프론트엔드","React 18, Vite 5, Zustand, React Router 6, Axios"],
  ["DB / 자동화","PostgreSQL 16, n8n(최신)"],
  ["빌드·배포","Docker Compose(4개 서비스), Nginx, GitHub Actions CI"],
  ["테스트","Jest (백엔드 단위·통합 테스트)"]],
 widths=[1.6,4.8])
h("4.2 주요 기능 구현 및 분석", 2)
bullets([
 ("인증·권한: ", "JWT(access 약 1시간 + refresh 7일) 기반. refresh 토큰은 HttpOnly 쿠키로만 전달하고, 로그아웃 시 tokenVersion을 증가시켜 발급된 토큰을 즉시 무효화한다. 역할별 가드와 회사코드 스코핑으로 접근을 통제한다."),
 ("퀘스트 워크플로우: ", "발행·착수·거부·재개봉·제출·검토 6개 상태 전이를 모두 '기대 상태 조건부 갱신'으로 원자화하여 동시성 문제를 차단한다."),
 ("CSV 일괄 발행/내보내기: ", "관리자가 CSV로 다건 퀘스트를 한 번에 발행(최대 100건, 단일 트랜잭션)하고 현황을 CSV로 내보낸다. 내보내기 시 '=,+,-,@'로 시작하는 셀을 무력화해 수식 인젝션을 방지한다."),
 ("증빙 파일: ", "PNG/JPEG/GIF/WEBP/PDF만 허용(최대 10MB)하며 DB에 BLOB로 저장한다. 응답에 nosniff를 적용하고 화이트리스트 형식만 inline 미리보기를 허용해 저장형 XSS를 차단한다."),
 ("실시간 알림: ", "발행/제출/검토/거부/재개봉/마감임박/마감초과 7개 이벤트를 n8n으로 전송한다. 사용자 응답을 막지 않도록 fire-and-forget로 호출한다."),
 ("마감 스케줄러: ", "매시간 크론으로 마감 임박(24h 이내)·마감 초과 퀘스트를 찾아 알림을 발송하며, 조건부 갱신으로 중복 알림을 방지한다."),
 ("감사 로그·통계: ", "권한 변경·검토·삭제 등 주요 행위를 감사 로그로 남기고, 진행률·담당자별 통계를 집계해 대시보드로 제공한다."),
])
para("[삽입 위치] 퀘스트 발행·제출·검토 화면 캡처, Slack 알림 도착 캡처, 진행률 대시보드 캡처", size=9.5, italic=True)
h("4.3 핵심 알고리즘 및 보안 설계", 2)
para("(1) 웹훅 HMAC 서명/검증 — 백엔드는 페이로드를 키 정렬 기반으로 정규 직렬화(stableStringify)한 뒤 공유 비밀키로 HMAC-SHA256 서명을 생성하여 헤더에 싣는다. n8n은 동일 규칙으로 재계산해 서명을 검증함으로써 위·변조되거나 비인가된 호출을 차단한다. 직렬화는 undefined 키를 생략해 항상 유효한 JSON을 보장하여 양측 결과의 바이트 일치를 보장한다.")
mono_block([
 "sig = HMAC_SHA256(secret, stableStringify({event, timestamp, data}))",
 "헤더: X-OnQuest-Signature, X-OnQuest-Timestamp, X-OnQuest-Event",
 "n8n: 타임스탬프 드리프트(±5분) 검사 + 서명 재계산 일치 시에만 처리",
])
para("(2) 증빙 외부 공유 토큰 — Slack에서 인증 없이 증빙을 열람할 수 있도록, 퀘스트 ID·만료시각을 포함한 HMAC 서명 토큰을 발급한다. 검증 시 길이 확인 후 timingSafeEqual로 비교하여 타이밍 공격을 방지하고, 만료된 링크는 거부한다.")
para("(3) 원자적 상태 전이 — '기대 상태인 행만' 갱신하는 조건부 쿼리로 한 행을 선점하고, 선점에 실패하면(이미 다른 요청이 상태를 바꾼 경우) 409 Conflict를 반환하여 중복 알림·이중 처리를 막는다.")
h("4.4 기존 SW 대비 차별성·창의성", 2)
bullets([
 ("온보딩+게임화 결합: ", "범용 업무 관리 도구나 정적 체크리스트와 달리, 신입 온보딩에 특화된 상태형 '퀘스트' 모델을 제시한다."),
 ("알림 자동화 내장: ", "별도 봇 개발 없이 n8n 워크플로우로 Slack 알림을 구성해, 코드 변경 없이 알림 시나리오를 확장할 수 있다."),
 ("보안 기본 탑재: ", "HMAC 서명, 토큰 폐기, 테넌트 격리, 감사 로그 등 실무 보안을 처음부터 설계에 반영했다."),
 ("셀프호스팅·저비용: ", "전 구성요소가 오픈소스이며 Docker 한 번으로 기동되어 도입 장벽이 낮다."),
])
pagebreak()

# ===== 5. 검증 =====
h("5. 검증 (기능·성능 평가)", 1)
h("5.1 기능 동작 체크리스트", 2)
table(["ID", "검증 항목", "기대 결과", "적/부"],
 [["T-01","회사 첫 가입자 슈퍼관리자 지정","최초 가입 시 role=superadmin","적"],
  ["T-02","단건/CSV 일괄 퀘스트 발행","퀘스트 생성 및 Slack 발행 알림","적"],
  ["T-03","신입 착수 → 증빙 제출","상태 검토대기 전환, 제출 알림","적"],
  ["T-04","관리자 반려(피드백)","상태 반려, 피드백 전달, 재제출 가능","적"],
  ["T-05","관리자 승인","상태 완료 처리","적"],
  ["T-06","마감 임박/초과 자동 알림","스케줄러가 중복 없이 알림 발송","적"],
  ["T-07","권한 외 접근 차단","역할/회사 불일치 시 권한 거부","적"],
  ["T-08","동시 검토 요청(경합)","한 건만 처리, 나머지 409 Conflict","적"],
  ["T-09","증빙 형식 검증","비허용 형식 업로드 거부","적"],
  ["T-10","웹훅 서명 검증","서명 불일치 페이로드 거부","적"]],
 widths=[0.6,2.3,2.7,0.7])
h("5.2 품질·성능 지표", 2)
table(["지표", "결과"],
 [["백엔드 자동화 테스트","6개 스위트 / 36개 테스트 전부 통과"],
  ["CI 파이프라인","GitHub Actions: 타입체크·린트·테스트·빌드 통과(green)"],
  ["Slack 알림 지연(목표)","이벤트 발생 후 3초 이내(웹훅 비동기, 2.5s 타임아웃)"],
  ["정적 분석","TypeScript 타입체크·ESLint 통과"],
  ["가용성","컨테이너 헬스체크 + DB 프로브 타임아웃(3s)으로 장애 감지"]],
 widths=[2.2,4.2])
para("※ 업무 효율 측면: 기존에는 사수가 신입에게 개별적으로 업무를 전달·확인했으나, On-Quest 도입 시 발행·알림·검토가 자동화되어 반복적인 전달·확인 커뮤니케이션이 감소하고 신입의 대기 시간이 줄어든다.")
pagebreak()

# ===== 6. 소감 =====
h("6. 소감 및 제언", 1)
para("이번 프로젝트를 통해 단일 기능 구현을 넘어 '서버–자동화 툴–메신저'를 안전하게 연결하는 시스템 통합 역량을 기를 수 있었다. 특히 웹훅에 HMAC 서명을 적용하고 양측 직렬화 결과의 일치를 보장하는 과정에서, 서드파티 연동 시 무결성·인증을 어떻게 보장하는지 실무적으로 체득했다.")
para("또한 동시성 환경에서의 상태 전이, 토큰 기반 인증의 폐기 전략, 멀티테넌시 데이터 격리 등 백엔드 개발자에게 필수적인 주제를 직접 설계·검증하며 실무 역량이 향상됨을 체감했다.")
para("제언: 향후 인증샷 트래픽이 증가할 경우 증빙 저장소를 외부 오브젝트 스토리지(S3 등)로 분리하고, 알림 채널을 Slack 외 이메일·문자로 확장하며, 온보딩 진척에 따른 배지·레벨 등 게임화 요소를 강화하면 활용도를 더 높일 수 있을 것이다.")
pagebreak()

# ===== 7. 참고사항 =====
h("7. 참고사항", 1)
h("7.1 용어/약어", 2)
table(["용어/약어", "설명"],
 [["퀘스트(Quest)","신입에게 부여되는 온보딩 과제 단위"],
  ["RBAC","Role-Based Access Control, 역할 기반 접근제어"],
  ["JWT","JSON Web Token, 토큰 기반 인증 방식"],
  ["HMAC","Hash-based Message Authentication Code, 메시지 무결성·인증 코드"],
  ["멀티테넌시","하나의 시스템이 다수 고객(회사) 데이터를 분리 운영하는 구조"],
  ["n8n","오픈소스 워크플로우 자동화 도구"],
  ["BLOB","Binary Large Object, 이진 대용량 데이터(여기선 증빙 파일)"]],
 widths=[1.6,4.8])
h("7.2 참고문헌", 2)
refs = ["NestJS 공식 문서, https://docs.nestjs.com","Prisma 공식 문서, https://www.prisma.io/docs",
 "React 공식 문서, https://react.dev","n8n 공식 문서, https://docs.n8n.io",
 "Slack API 문서, https://api.slack.com","PostgreSQL 16 Documentation, https://www.postgresql.org/docs/16/",
 "OWASP, CSV Injection / Formula Injection, https://owasp.org"]
for i, rf in enumerate(refs, 1):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"[{i}] {rf}"); kfont_run(r, BODY, 10)
pagebreak()

# ===== 부록 =====
h("부록. 실행 방법 및 소스 구조", 1)
h("A. 실행 방법 (Docker Compose)", 2)
mono_block([
 "# 1) 환경 변수 준비",
 "cp .env.example .env   # JWT_SECRET 등 시크릿 값 설정(필수)",
 "",
 "# 2) 전체 스택 기동 (DB·n8n·백엔드·프론트)",
 "docker compose up -d --build",
 "",
 "# 3) 접속",
 "Frontend : http://localhost:8080",
 "Backend  : http://localhost:3000/api  (health: /api/health)",
 "n8n      : http://localhost:5678",
])
para("※ 로컬 개발 시 backend/frontend 각각 npm install 후 npm run start:dev / npm run dev 로 구동 가능.", size=9.5, italic=True)
h("B. 소스 구조 (요약)", 2)
mono_block([
 "on-quest/",
 "├─ backend/        # NestJS API",
 "│  ├─ src/auth     # 인증·JWT·권한",
 "│  ├─ src/quest    # 퀘스트·증빙·검토·CSV",
 "│  ├─ src/automation  # n8n 웹훅(HMAC)",
 "│  ├─ src/audit    # 감사 로그",
 "│  ├─ src/health   # 헬스체크",
 "│  └─ prisma/      # 스키마·마이그레이션",
 "├─ frontend/       # React + Vite",
 "│  └─ src/{auth, components, pages, store, api, utils}",
 "├─ n8n/            # 워크플로우 템플릿",
 "├─ docker-compose.yml",
 "└─ .github/workflows/ci.yml",
])
para("[삽입 위치] 핵심 소스코드 발췌 또는 GitHub 링크", size=9.5, italic=True)

doc.save(r"D:\최종보고서.docx")
print("SAVED")
